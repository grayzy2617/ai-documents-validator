import sys
import os
import docx

sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from main import app
from fastapi.testclient import TestClient

client = TestClient(app)

def create_bad_docx(file_path):
    doc = docx.Document()
    doc.add_paragraph('cộng hòa xã hội chủ nghĩa việt nam')
    doc.add_paragraph('Độc Lập - tự do - Hạnh phúc')
    doc.save(file_path)

def run_tests():
    print("=== TEST AUTOFIX ===")
    
    # 1. Login
    res = client.post("/login", data={"username": "admin", "password": "123456"})
    assert res.status_code == 200, res.text
    token = res.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    
    # 2. Upload file
    test_file_path = os.path.join(os.path.dirname(__file__), "bad_doc.docx")
    create_bad_docx(test_file_path)
    
    with open(test_file_path, "rb") as f:
        res = client.post("/documents/check", headers=headers, files={"file": ("bad_doc.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    
    assert res.status_code == 200, f"Upload lỗi: {res.text}"
    doc_id = res.json()["document_id"]
    print(f"Đã upload và kiểm tra AI. Document ID: {doc_id}")
    
    # 3. Trigger autofix
    print("Đang gọi AI sinh Autofix Plan và sửa file...")
    res = client.post(f"/documents/{doc_id}/autofix", headers=headers)
    assert res.status_code == 200, f"Lỗi Autofix: {res.text}"
    print(f"Autofix Response: {res.json()}")
    
    # 4. Download autofixed file
    res = client.get(f"/documents/{doc_id}/autofixed_file", headers=headers)
    assert res.status_code == 200, f"Lỗi Download Autofixed: {res.text}"
    
    with open("downloaded_autofixed.docx", "wb") as f:
        f.write(res.content)
    print("Đã tải xuống file autofixed thành công!")
    
    # Clean up
    if os.path.exists(test_file_path): os.remove(test_file_path)
    if os.path.exists("downloaded_autofixed.docx"): os.remove("downloaded_autofixed.docx")

if __name__ == "__main__":
    run_tests()
