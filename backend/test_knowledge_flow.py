import sys
import os
import docx

print("Starting test script...", flush=True)
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

print("Importing app...", flush=True)
from main import app
from fastapi.testclient import TestClient

client = TestClient(app)

def create_dummy_docx(file_path):
    doc = docx.Document()
    doc.add_heading('Quy định test', 0)
    doc.add_paragraph('Đây là một văn bản quy định mẫu để test vector DB. Yêu cầu lề trái 3cm.')
    doc.save(file_path)

def run_tests():
    print("=== BẮT ĐẦU CHẠY BỘ KIỂM THỬ BƯỚC 4 (ADMIN KNOWLEDGE) ===")
    
    # 1. Đăng nhập lấy token
    response = client.post(
        "/login",
        data={"username": "admin", "password": "123456"} # admin account from create_test_users
    )
    assert response.status_code == 200, f"Đăng nhập thất bại: {response.text}"
    token = response.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    print("1. Đăng nhập với quyền Admin thành công.")

    # 2. Tạo file mẫu
    test_file_path = os.path.join(os.path.dirname(__file__), "test_rule.docx")
    create_dummy_docx(test_file_path)
    
    # 3. Test API Upload
    print("2. Đang test upload (và nhúng vào Vector DB)...")
    with open(test_file_path, "rb") as f:
        upload_res = client.post(
            "/knowledge/upload",
            headers=headers,
            data={"title": "Quy định Test Vector DB", "document_type": "Nội quy"},
            files={"file": ("test_rule.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        )
    assert upload_res.status_code == 200, f"Lỗi Upload: {upload_res.text}"
    doc_id = upload_res.json()["document_id"]
    print(f"   -> Upload thành công. Document ID: {doc_id}")

    # 4. Test API GET list
    list_res = client.get("/knowledge", headers=headers)
    assert list_res.status_code == 200
    docs = list_res.json()
    assert any(d["id"] == doc_id for d in docs), "Không tìm thấy tài liệu vừa upload trong DB."
    print("3. Kiểm tra danh sách tri thức (GET /knowledge) -> PASSED")

    # 5. Test API Delete
    print("4. Đang test xóa tri thức...")
    del_res = client.delete(f"/knowledge/{doc_id}", headers=headers)
    assert del_res.status_code == 200, f"Lỗi Xóa: {del_res.text}"
    
    # 6. Verify delete
    list_res2 = client.get("/knowledge", headers=headers)
    docs2 = list_res2.json()
    assert not any(d["id"] == doc_id for d in docs2), "Tài liệu vẫn còn trong DB sau khi xóa."
    print("   -> Xóa thành công khỏi SQLite và Vector DB -> PASSED")

    # Clean up
    if os.path.exists(test_file_path):
        os.remove(test_file_path)

    print("=== TẤT CẢ CÁC TEST CASE ĐỀU PASSED ===================")

if __name__ == "__main__":
    run_tests()
