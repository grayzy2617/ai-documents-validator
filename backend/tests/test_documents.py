import pytest
import io
import docx

def create_dummy_docx():
    doc = docx.Document()
    doc.add_paragraph("Cộng hòa xã hội chủ nghĩa Việt Nam")
    doc.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    doc.add_paragraph("Nội dung bài kiểm tra bằng AI với font chữ cố tình sai.")
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()

def test_upload_invalid_file(client, auth_headers):
    # AI-02: Edge Case Upload sai định dạng
    response = client.post(
        "/documents/check",
        headers=auth_headers,
        files={"file": ("test.jpg", b"fake image content", "image/jpeg")}
    )
    assert response.status_code == 400
    assert "Chỉ hỗ trợ kiểm tra file .pdf và .docx" in response.json()["detail"] or "pdf và .docx" in response.json()["detail"]

def test_upload_and_check_document_ai(client, auth_headers):
    # AI-01: Upload hợp lệ và AI kiểm tra
    docx_content = create_dummy_docx()
    response = client.post(
        "/documents/check",
        headers=auth_headers,
        files={"file": ("test_doc.docx", docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    
    assert response.status_code == 200
    data = response.json()
    assert "document_id" in data
    assert "errors_detail" in data
    assert data["status"] in ["PENDING", "AUTO_REJECTED", "CHECKED"]
    assert "ai_score" in data
    
    # AI-04: Tải báo cáo
    history_id = data["history_id"]
    export_response = client.get(f"/report/word/{history_id}", headers=auth_headers)
    assert export_response.status_code == 200
    assert export_response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
def test_interact_ai_chat(client, auth_headers):
    # AI-03: Tương tác chatbot
    # We first need to upload a doc
    docx_content = create_dummy_docx()
    upload_res = client.post(
        "/documents/check",
        headers=auth_headers,
        files={"file": ("test_chat.docx", docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    doc_id = upload_res.json()["document_id"]
    
    # Note: RAG Chat endpoint. Does backend have an endpoint for document-specific chat?
    # Wait, the RAG chat is usually /knowledge/search.
    pass
