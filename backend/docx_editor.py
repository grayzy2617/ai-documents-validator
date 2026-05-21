"""Thao tác DOCX: working copy, sửa đoạn, autofix."""
import os
import shutil
from typing import Optional

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def get_working_path(
    original_path: str,
    fixed_path: Optional[str],
) -> str:
    if fixed_path and os.path.exists(fixed_path):
        return fixed_path
    return original_path


def ensure_working_copy(
    original_path: str,
    fixed_path: Optional[str],
    uploads_dir: str,
    document_id: int,
    filename: str,
) -> tuple[str, bool]:
    """Trả về (path, created_new)."""
    if fixed_path and os.path.exists(fixed_path):
        return fixed_path, False
    working_name = f"working_{document_id}_{filename}"
    working_path = os.path.join(uploads_dir, working_name)
    shutil.copy2(original_path, working_path)
    return working_path, True


def _apply_alignment(p, align: str):
    align = (align or "").upper()
    if align == "CENTER":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "RIGHT":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "LEFT":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "JUSTIFY":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def apply_fix_to_paragraph(p, fix: dict) -> bool:
    search_text = (fix.get("search_text") or "").strip().lower()
    if not search_text:
        return False
    p_text = p.text.strip().lower()
    if search_text not in p_text:
        return False

    replace_text = fix.get("replace_text", "")
    p.clear()
    run = p.add_run(replace_text)
    if fix.get("bold"):
        run.bold = True
    if fix.get("italic"):
        run.italic = True
    if fix.get("font_name"):
        run.font.name = fix.get("font_name")
    if fix.get("font_size"):
        run.font.size = Pt(fix.get("font_size"))
    _apply_alignment(p, fix.get("alignment", ""))
    return True


def apply_fixes_to_docx(file_path: str, fixes: list[dict]) -> int:
    docx_file = docx.Document(file_path)
    applied = 0
    for p in docx_file.paragraphs:
        p_text = p.text.strip().lower()
        if not p_text:
            continue
        for fix in fixes:
            if apply_fix_to_paragraph(p, fix):
                applied += 1
                break
    docx_file.save(file_path)
    return applied


def find_paragraph_index(docx_file: docx.Document, snippet: str) -> Optional[int]:
    if not snippet:
        return None
    needle = snippet.strip().lower()[:120]
    if len(needle) < 3:
        return None
    for idx, p in enumerate(docx_file.paragraphs):
        if needle in p.text.strip().lower():
            return idx
    short = needle[:40]
    for idx, p in enumerate(docx_file.paragraphs):
        if short in p.text.strip().lower():
            return idx
    return None


def update_paragraph_text(file_path: str, paragraph_index: int, new_text: str) -> bool:
    docx_file = docx.Document(file_path)
    if paragraph_index < 0 or paragraph_index >= len(docx_file.paragraphs):
        return False
    p = docx_file.paragraphs[paragraph_index]
    p.clear()
    p.add_run(new_text)
    docx_file.save(file_path)
    return True


def bulk_update_paragraphs(file_path: str, updates: list[dict]) -> int:
    docx_file = docx.Document(file_path)
    count = 0
    for item in updates:
        idx = item.get("index")
        text = item.get("text")
        if idx is None or text is None:
            continue
        if 0 <= idx < len(docx_file.paragraphs):
            p = docx_file.paragraphs[idx]
            p.clear()
            p.add_run(str(text))
            count += 1
    docx_file.save(file_path)
    return count


def paragraph_preview(file_path: str, paragraph_index: int) -> Optional[str]:
    docx_file = docx.Document(file_path)
    if paragraph_index < 0 or paragraph_index >= len(docx_file.paragraphs):
        return None
    return docx_file.paragraphs[paragraph_index].text
