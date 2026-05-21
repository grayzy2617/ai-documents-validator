from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query
from sqlalchemy.orm import Session
from sqlalchemy import desc
from pydantic import BaseModel
import os
import shutil
from fastapi.responses import FileResponse

import models, schemas
from database import get_db
from routers.auth import get_current_user
from document_processor import extract_text, extract_plain_text, docx_to_html
from docx_editor import (
    ensure_working_copy,
    get_working_path,
    apply_fixes_to_docx,
    find_paragraph_index,
    update_paragraph_text,
    bulk_update_paragraphs,
    paragraph_preview,
)
from llm_service import check_document_errors, generate_autofix_plan, summarize_document, score_document_structure
import docx

router = APIRouter()

# Thư mục lưu trữ văn bản của người dùng đẩy lên
USER_UPLOADS_DIR = os.path.join(os.path.dirname(__file__), "..", "user_uploads")
os.makedirs(USER_UPLOADS_DIR, exist_ok=True)


def _user_roles(user: models.User) -> list:
    return [r.role_name for r in user.roles]


def _get_document_for_access(document_id: int, current_user: models.User, db: Session) -> models.UserDocument:
    doc = db.query(models.UserDocument).filter(models.UserDocument.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Không tìm thấy tài liệu")
    roles = _user_roles(current_user)
    if doc.user_id != current_user.id and "BGH" not in roles and "TO_TRUONG" not in roles:
        raise HTTPException(status_code=403, detail="Không có quyền truy cập tài liệu này")
    return doc


def _latest_check_history(doc: models.UserDocument, db: Session) -> models.CheckHistory | None:
    return (
        db.query(models.CheckHistory)
        .filter(models.CheckHistory.document_id == doc.id)
        .order_by(desc(models.CheckHistory.check_date))
        .first()
    )


def _require_docx(doc: models.UserDocument):
    if not doc.original_file_name.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Tính năng chỉ hỗ trợ file .docx")


def _get_or_create_working_path(
    doc: models.UserDocument,
    history: models.CheckHistory,
    db: Session,
) -> str:
    _require_docx(doc)
    path, created = ensure_working_copy(
        doc.file_path,
        history.fixed_file_path,
        USER_UPLOADS_DIR,
        doc.id,
        doc.original_file_name,
    )
    if created or history.fixed_file_path != path:
        history.fixed_file_path = path
        db.commit()
    return path


def _get_error_for_owner(error_id: int, current_user: models.User, db: Session) -> models.DetectedError:
    err = db.query(models.DetectedError).filter(models.DetectedError.id == error_id).first()
    if not err:
        raise HTTPException(status_code=404, detail="Không tìm thấy lỗi")
    if err.history.document.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Không có quyền thao tác trên lỗi này")
    return err

@router.post("/check")
async def check_user_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """
    API nhận file từ người dùng, đọc chữ, đẩy qua AI bắt lỗi và lưu vào DB.
    """
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in [".pdf", ".docx"]:
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ kiểm tra file .pdf và .docx")
    
    # Đặt tên file chống trùng lặp theo ID user
    saved_filename = f"{current_user.id}_{file.filename}"
    file_path = os.path.join(USER_UPLOADS_DIR, saved_filename)
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    try:
        # 1. Trích xuất text từ file
        text = extract_text(file_path)
        if not text.strip():
            raise HTTPException(status_code=400, detail="Không thấy văn bản hoặc file rỗng")
        
        # Lấy metadata dung lượng file
        metadata_info = {"file_size_bytes": os.path.getsize(file_path)}
        
        # 2. Tạo bản ghi UserDocument
        new_doc = models.UserDocument(
            user_id=current_user.id,
            original_file_name=file.filename,
            file_path=file_path,
            status="CHECKED",
            metadata_info=metadata_info
        )
        db.add(new_doc)
        db.flush() # Lấy new_doc.id ngay lập tức
        
        # 3. Tạo bản ghi lịch sử kiểm tra (CheckHistory)
        new_history = models.CheckHistory(
            document_id=new_doc.id
        )
        db.add(new_history)
        db.flush() # Lấy new_history.id
        
        # 4. GỌI BỘ NÃO AI (Luồng RAG Core)
        detected_errors_list = check_document_errors(text)
        
        # 5. Lưu danh sách lỗi vào Database (DetectedError)
        saved_errors = []
        if isinstance(detected_errors_list, list):
            for err in detected_errors_list:
                db_error = models.DetectedError(
                    history_id=new_history.id,
                    error_type=err.get("error_type", "Lỗi không xác định"),
                    error_location=err.get("error_location", "Không rõ vị trí"),
                    description=err.get("description", ""),
                    suggestion=err.get("suggestion", ""),
                    status="UNFIXED"
                )
                db.add(db_error)
                saved_errors.append(db_error)
                
        db.commit()
        
        # 6. Chấm điểm cấu trúc AI (v2.0)
        try:
            ai_score = score_document_structure(text, len(saved_errors))
            new_doc.ai_score = ai_score
            
            # Nếu score < 80 → Auto-reject
            if ai_score < 80:
                new_doc.status = "AUTO_REJECTED"
            else:
                new_doc.status = "PENDING"  # Chờ Tổ trưởng duyệt
                
            db.commit()
        except Exception as score_err:
            print(f"Lỗi chấm điểm AI: {score_err}")
            new_doc.ai_score = None
            new_doc.status = "CHECKED"
            db.commit()
        
        return {
            "message": "Đã kiểm tra thành công",
            "document_id": new_doc.id,
            "history_id": new_history.id,
            "total_errors": len(saved_errors),
            "errors_detail": detected_errors_list,
            "ai_score": new_doc.ai_score,
            "status": new_doc.status
        }
        
    except Exception as e:
        db.rollback()
        # Xóa file tạm nếu lỗi
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"Lỗi xử lý kiểm tra: {str(e)}")

@router.get("/history")
def get_user_documents_history(
    page: int = 1,
    size: int = 10,
    db: Session = Depends(get_db), 
    current_user: models.User = Depends(get_current_user)
):
    """API lấy ra danh sách các file văn bản user đã tải lên kiểm tra có phân trang"""
    offset = (page - 1) * size
    docs = db.query(models.UserDocument).filter(
        models.UserDocument.user_id == current_user.id
    ).order_by(
        models.UserDocument.created_at.desc()
    ).offset(offset).limit(size).all()
    
    total = db.query(models.UserDocument).filter(
        models.UserDocument.user_id == current_user.id
    ).count()

    return {
        "items": docs,
        "total": total,
        "page": page,
        "size": size,
        "pages": (total + size - 1) // size
    }

@router.get("/{document_id}/meta")
def get_document_meta(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Metadata văn bản (dùng cho trang kết quả / xem trước)."""
    doc = _get_document_for_access(document_id, current_user, db)
    history = _latest_check_history(doc, db)
    ext = os.path.splitext(doc.original_file_name)[1].lower()
    has_fixed = bool(
        history
        and history.fixed_file_path
        and os.path.exists(history.fixed_file_path)
    )
    return {
        "document_id": doc.id,
        "history_id": history.id if history else None,
        "original_file_name": doc.original_file_name,
        "file_type": "pdf" if ext == ".pdf" else "docx" if ext == ".docx" else ext,
        "status": doc.status,
        "has_fixed_file": has_fixed,
    }


@router.get("/{document_id}/file")
async def get_document_file(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Trả về file gốc để xem trước (PDF trong react-pdf)."""
    doc = _get_document_for_access(document_id, current_user, db)
    if not os.path.exists(doc.file_path):
        raise HTTPException(status_code=404, detail="File không tồn tại trên hệ thống")
    media = "application/pdf" if doc.original_file_name.lower().endswith(".pdf") else (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    return FileResponse(doc.file_path, filename=doc.original_file_name, media_type=media)


@router.get("/{document_id}/preview-text")
def get_document_preview_text(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Nội dung text thuần (fallback)."""
    doc = _get_document_for_access(document_id, current_user, db)
    text = extract_plain_text(doc.file_path)
    return {
        "text": text[:80000],
        "original_file_name": doc.original_file_name,
    }


@router.get("/{document_id}/preview-html")
def get_document_preview_html(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """HTML render từ DOCX để hiển thị giống Word."""
    doc = _get_document_for_access(document_id, current_user, db)
    if not doc.original_file_name.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Preview HTML chỉ hỗ trợ file .docx")
    history = _latest_check_history(doc, db)
    preview_path = doc.file_path
    if (
        history
        and history.fixed_file_path
        and os.path.exists(history.fixed_file_path)
    ):
        preview_path = history.fixed_file_path
    return {
        "html": docx_to_html(preview_path),
        "original_file_name": doc.original_file_name,
        "is_fixed_version": preview_path != doc.file_path,
    }


@router.get("/{document_id}/errors")
def get_errors_by_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Danh sách lỗi theo document_id (lịch sử kiểm tra mới nhất)."""
    doc = _get_document_for_access(document_id, current_user, db)
    history = _latest_check_history(doc, db)
    if not history:
        return []
    return db.query(models.DetectedError).filter(models.DetectedError.history_id == history.id).all()


@router.get("/history/{history_id}/errors")
def get_errors_by_check_history(history_id: int, db: Session = Depends(get_db), current_user: models.User = Depends(get_current_user)):
    """API xem chi tiết các lỗi theo history_id (tương thích cũ)."""
    history = db.query(models.CheckHistory).filter(models.CheckHistory.id == history_id).first()
    if not history:
        raise HTTPException(status_code=404, detail="Không tìm thấy lịch sử")
    roles = _user_roles(current_user)
    if history.document.user_id != current_user.id and "BGH" not in roles and "TO_TRUONG" not in roles:
        raise HTTPException(status_code=403, detail="Không có quyền truy cập hoặc không tìm thấy lịch sử")

    errors = db.query(models.DetectedError).filter(models.DetectedError.history_id == history_id).all()
    return errors

class ErrorFixStatus(BaseModel):
    status: str # Ví dụ: MANUAL_FIXED, IGNORED

@router.put("/errors/{error_id}/fix")
def update_error_fix_status(
    error_id: int, 
    status_data: ErrorFixStatus,
    db: Session = Depends(get_db), 
    current_user: models.User = Depends(get_current_user)
):
    """API cho phép người dùng tự đánh dấu một lỗi là đã sửa hoặc bỏ qua"""
    err = db.query(models.DetectedError).filter(models.DetectedError.id == error_id).first()
    if not err:
        raise HTTPException(status_code=404, detail="Không tìm thấy lỗi")
        
    # Xác minh user là chủ sở hữu của văn bản
    if err.history.document.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Không có quyền thao tác trên lỗi này")
        
    err.status = status_data.status
    db.commit()
    db.refresh(err)
    return err

@router.post("/{document_id}/autofix")
def trigger_autofix(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """Autofix tất cả lỗi chưa xử lý trên bản working copy."""
    doc = _get_document_for_access(document_id, current_user, db)
    if doc.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Chỉ chủ văn bản mới được sửa")
    _require_docx(doc)

    history = _latest_check_history(doc, db)
    if not history:
        raise HTTPException(status_code=400, detail="Văn bản chưa được kiểm tra")

    errors = [e for e in history.errors if e.status not in ["IGNORED", "MANUAL_FIXED", "AUTO_FIXED"]]
    if not errors:
        return {"message": "Không có lỗi nào cần Autofix", "fixed_file_path": history.fixed_file_path}

    working_path = _get_or_create_working_path(doc, history, db)
    text = extract_text(working_path)
    plan = generate_autofix_plan(text, errors)
    if not plan:
        raise HTTPException(status_code=500, detail="AI không thể tạo được kế hoạch sửa lỗi")

    apply_fixes_to_docx(working_path, plan)
    for e in errors:
        e.status = "AUTO_FIXED"
    db.commit()

    return {
        "message": "Đã tự động chuẩn hóa văn bản",
        "fixed_file_path": working_path,
        "applied_count": len(plan),
    }


class AiFixApplyBody(BaseModel):
    search_text: str | None = None
    replace_text: str | None = None
    bold: bool | None = None
    italic: bool | None = None
    alignment: str | None = None


@router.post("/{document_id}/errors/{error_id}/ai-fix/preview")
def preview_ai_fix_single_error(
    document_id: int,
    error_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Xem diff trước khi áp dụng sửa AI cho một lỗi."""
    doc = _get_document_for_access(document_id, current_user, db)
    if doc.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Chỉ chủ văn bản mới được sửa")
    _require_docx(doc)

    err = _get_error_for_owner(error_id, current_user, db)
    if err.history.document_id != document_id:
        raise HTTPException(status_code=400, detail="Lỗi không thuộc văn bản này")

    history = _latest_check_history(doc, db)
    working_path = _get_or_create_working_path(doc, history, db)
    text = extract_text(working_path)
    plan = generate_autofix_plan(text, [err])
    if not plan:
        raise HTTPException(status_code=500, detail="AI không tạo được gợi ý sửa cho lỗi này")

    fix = plan[0]
    docx_file = docx.Document(working_path)
    para_idx = find_paragraph_index(
        docx_file,
        fix.get("search_text") or err.error_location or err.description,
    )
    before = docx_file.paragraphs[para_idx].text if para_idx is not None else ""
    after = before
    search = fix.get("search_text", "")
    replace = fix.get("replace_text", "")
    if search and search in before:
        after = before.replace(search, replace, 1)
    elif replace:
        after = replace

    return {
        "error_id": error_id,
        "paragraph_index": para_idx,
        "before_text": before,
        "after_text": after,
        "search_text": search,
        "replace_text": replace,
        "fix": fix,
    }


@router.post("/{document_id}/errors/{error_id}/ai-fix/apply")
def apply_ai_fix_single_error(
    document_id: int,
    error_id: int,
    body: AiFixApplyBody | None = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Áp dụng sửa AI cho một lỗi (sau khi xem preview)."""
    doc = _get_document_for_access(document_id, current_user, db)
    if doc.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Chỉ chủ văn bản mới được sửa")
    _require_docx(doc)

    err = _get_error_for_owner(error_id, current_user, db)
    if err.history.document_id != document_id:
        raise HTTPException(status_code=400, detail="Lỗi không thuộc văn bản này")

    history = _latest_check_history(doc, db)
    working_path = _get_or_create_working_path(doc, history, db)

    fix = None
    if body and body.search_text and body.replace_text is not None:
        fix = {
            "search_text": body.search_text,
            "replace_text": body.replace_text,
            "bold": body.bold,
            "italic": body.italic,
            "alignment": body.alignment,
        }
    else:
        text = extract_text(working_path)
        plan = generate_autofix_plan(text, [err])
        if not plan:
            raise HTTPException(status_code=500, detail="AI không tạo được kế hoạch sửa")
        fix = plan[0]

    applied = apply_fixes_to_docx(working_path, [fix])
    if applied == 0 and fix.get("replace_text") is not None:
        snippet = err.error_location or err.description or ""
        docx_file = docx.Document(working_path)
        para_idx = find_paragraph_index(docx_file, snippet)
        if para_idx is not None:
            update_paragraph_text(working_path, para_idx, fix["replace_text"])
            applied = 1

    if applied == 0:
        raise HTTPException(status_code=400, detail="Không tìm thấy đoạn văn bản để áp dụng sửa")

    err.status = "AUTO_FIXED"
    db.commit()

    return {
        "message": "Đã áp dụng sửa AI",
        "error_id": error_id,
        "has_fixed_file": True,
    }


class ParagraphItem(BaseModel):
    index: int
    text: str


class ParagraphBulkBody(BaseModel):
    paragraphs: list[ParagraphItem]


@router.put("/{document_id}/paragraphs")
def save_paragraph_edits(
    document_id: int,
    body: ParagraphBulkBody,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Lưu chỉnh sửa thủ công từ editor (WYSIWYG)."""
    doc = _get_document_for_access(document_id, current_user, db)
    if doc.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Chỉ chủ văn bản mới được sửa")
    _require_docx(doc)

    history = _latest_check_history(doc, db)
    if not history:
        raise HTTPException(status_code=400, detail="Văn bản chưa được kiểm tra")

    working_path = _get_or_create_working_path(doc, history, db)
    updates = [{"index": p.index, "text": p.text} for p in body.paragraphs]
    count = bulk_update_paragraphs(working_path, updates)
    if count == 0:
        raise HTTPException(status_code=400, detail="Không có đoạn nào được cập nhật")

    db.commit()
    return {
        "message": f"Đã lưu {count} đoạn văn bản",
        "updated_count": count,
        "has_fixed_file": True,
    }


@router.post("/{document_id}/recheck")
def recheck_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Kiểm tra lại văn bản (ưu tiên bản working / đã sửa)."""
    doc = _get_document_for_access(document_id, current_user, db)
    if doc.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Chỉ chủ văn bản mới được kiểm tra lại")

    history = _latest_check_history(doc, db)
    if not history:
        raise HTTPException(status_code=400, detail="Văn bản chưa được kiểm tra lần đầu")

    if doc.original_file_name.lower().endswith(".docx"):
        check_path = get_working_path(doc.file_path, history.fixed_file_path)
        if not os.path.exists(check_path):
            check_path = doc.file_path
    else:
        check_path = doc.file_path

    text = extract_text(check_path)
    if not text.strip():
        raise HTTPException(status_code=400, detail="Không trích xuất được nội dung để kiểm tra lại")

    detected = check_document_errors(text)
    new_history = models.CheckHistory(
        document_id=doc.id,
        fixed_file_path=history.fixed_file_path if history.fixed_file_path else None,
    )
    db.add(new_history)
    db.flush()

    saved = []
    if isinstance(detected, list):
        for item in detected:
            db_error = models.DetectedError(
                history_id=new_history.id,
                error_type=item.get("error_type", "Lỗi không xác định"),
                error_location=item.get("error_location", "Không rõ vị trí"),
                description=item.get("description", ""),
                suggestion=item.get("suggestion", ""),
                status="UNFIXED",
            )
            db.add(db_error)
            saved.append(db_error)

    try:
        doc.ai_score = score_document_structure(text, len(saved))
        if doc.ai_score < 80:
            doc.status = "AUTO_REJECTED"
        else:
            doc.status = "PENDING"
    except Exception:
        doc.status = "CHECKED"

    db.commit()

    return {
        "message": "Đã kiểm tra lại văn bản",
        "history_id": new_history.id,
        "total_errors": len(saved),
        "ai_score": doc.ai_score,
        "status": doc.status,
    }

@router.get("/{document_id}/autofixed_file")
async def get_autofixed_file(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Tải xuống file đã được chuẩn hóa."""
    doc = _get_document_for_access(document_id, current_user, db)
    history = _latest_check_history(doc, db)
    if not history or not history.fixed_file_path or not os.path.exists(history.fixed_file_path):
        raise HTTPException(status_code=404, detail="Không tìm thấy file đã chuẩn hóa")
    return FileResponse(
        history.fixed_file_path,
        filename=f"ChuanHoa_{doc.original_file_name}",
    )

@router.get("/{document_id}/summarize")
def summarize_user_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """API tóm tắt văn bản bằng AI (v2.0)"""
    doc = db.query(models.UserDocument).filter(models.UserDocument.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Không tìm thấy tài liệu")
    
    _get_document_for_access(document_id, current_user, db)

    text = extract_plain_text(doc.file_path)
    if not text.strip():
        raise HTTPException(status_code=400, detail="Không thể trích xuất nội dung văn bản")

    try:
        return summarize_document(text)
    except Exception as e:
        raise HTTPException(
            status_code=502,
            detail=f"Lỗi dịch vụ AI tóm tắt: {str(e)}",
        ) from e
