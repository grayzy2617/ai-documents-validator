import os
from io import BytesIO
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt, Inches

import models
from database import get_db
from routers.auth import get_current_user

router = APIRouter()

@router.get("/word/{history_id}")
def export_word_report(
    history_id: int, 
    db: Session = Depends(get_db), 
    current_user: models.User = Depends(get_current_user)
):
    """
    API xuất Báo cáo lỗi chi tiết ra file Word (.docx)
    """
    # 1. Truy vấn thông tin quá trình kiểm tra (Lịch sử, Document, Danh sách lỗi)
    history = db.query(models.CheckHistory).filter(models.CheckHistory.id == history_id).first()
    
    if not history:
        raise HTTPException(status_code=404, detail="Không tìm thấy lịch sử kiểm tra này")
        
    # Check quyền truy cập (chỉ được export file của chính mình hoặc admin)
    if history.document.user_id != current_user.id: # Tạm thời chưa có role Admin, nên gán chỉ chủ sở hữu
        raise HTTPException(status_code=403, detail="Không có quyền truy cập xuất báo cáo này")
        
    doc_info = history.document
    errors = history.errors
    
    # 2. Khởi tạo đối tượng Word Document
    doc = Document()
    
    # 3. Tạo Tiêu đề Báo Cáo
    title = doc.add_heading('BÁO CÁO KIỂM TRA THỂ THỨC VĂN BẢN', level=0)
    title.alignment = 1  # Canh giữa
    
    # 4. Tạo phần Thông tin chung (Metadata)
    doc.add_paragraph(f"Tệp tin gốc: {doc_info.original_file_name}", style='List Bullet')
    doc.add_paragraph(f"Người kiểm tra: {current_user.full_name} ({current_user.email})", style='List Bullet')
    
    check_time = history.check_date.strftime('%d/%m/%Y %H:%M:%S') if history.check_date else 'Không xác định'
    doc.add_paragraph(f"Thời gian kiểm tra: {check_time}", style='List Bullet')
    
    doc.add_paragraph(f"Tổng số lỗi phát hiện: {len(errors)} lỗi", style='List Bullet')
    
    doc.add_heading('CHI TIẾT LỖI PHÁT HIỆN:', level=1)
    
    # 5. Đổ dữ liệu vào bảng (Table)
    if not errors:
        doc.add_paragraph("Chúc mừng! Hệ thống AI không phát hiện vi phạm thể thức nào trong văn bản của bạn.")
    else:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Đặt tiêu đề cho các cột
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'STT'
        hdr_cells[1].text = 'Loại lỗi'
        hdr_cells[2].text = 'Vị trí & Mô tả vi phạm'
        hdr_cells[3].text = 'Gợi ý từ hệ thống AI'
        
        # Thiết lập độ rộng cột (tương đối)
        hdr_cells[0].width = Inches(0.5)
        hdr_cells[1].width = Inches(1.5)
        hdr_cells[2].width = Inches(2.5)
        hdr_cells[3].width = Inches(2.5)

        for index, error in enumerate(errors, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index)
            row_cells[1].text = error.error_type or "Không xác định"
            
            location = error.error_location or "Toàn văn bản"
            description = error.description or "Không có mô tả chi tiết"
            row_cells[2].text = f"[Vị trí: {location}]\n{description}"
            
            row_cells[3].text = error.suggestion or ""

    # 6. Chuẩn bị luồng Stream file trả về HTTP Response
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Tạo tên file báo cáo
    download_filename = f"Bao_cao_loi_{doc_info.id}.docx"
    
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={download_filename}"}
    )

from pydantic import BaseModel
from typing import List

class MergeReportRequest(BaseModel):
    history_ids: List[int]

@router.post("/merge")
def merge_reports(
    request: MergeReportRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """
    API tự động gom nhiều báo cáo kiểm tra thành 1 báo cáo tổng
    """
    roles = [r.role_name for r in current_user.roles]
    if "TO_TRUONG" not in roles and "BGH" not in roles:
        raise HTTPException(status_code=403, detail="Chỉ Tổ trưởng và BGH mới có quyền gom báo cáo")
        
    if not request.history_ids:
        raise HTTPException(status_code=400, detail="Không có báo cáo nào để gộp")
        
    doc = Document()
    title = doc.add_heading('BÁO CÁO TỔNG HỢP KIỂM TRA THỂ THỨC VĂN BẢN', level=0)
    title.alignment = 1
    
    doc.add_paragraph(f"Người xuất báo cáo: {current_user.full_name}")
    from datetime import datetime
    doc.add_paragraph(f"Thời gian xuất: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    
    for hid in request.history_ids:
        history = db.query(models.CheckHistory).filter(models.CheckHistory.id == hid).first()
        if not history:
            continue
            
        doc_info = history.document
        errors = history.errors
        
        doc.add_heading(f"Tài liệu: {doc_info.original_file_name}", level=1)
        doc.add_paragraph(f"Người nộp: {doc_info.owner} (ID: {doc_info.user_id})")
        doc.add_paragraph(f"Tổng số lỗi: {len(errors)}")
        
        if not errors:
            doc.add_paragraph("Không có lỗi nào được phát hiện.")
        else:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'STT'
            hdr_cells[1].text = 'Loại lỗi'
            hdr_cells[2].text = 'Mô tả'
            hdr_cells[3].text = 'Gợi ý'
            
            for index, error in enumerate(errors, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(index)
                row_cells[1].text = error.error_type or ""
                row_cells[2].text = error.description or ""
                row_cells[3].text = error.suggestion or ""
                
        doc.add_page_break()

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=Bao_cao_tong_hop.docx"}
    )
