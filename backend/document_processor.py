import os
import html
import fitz  # PyMuPDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ocr_service import needs_ocr_fallback, ocr_pdf_pages


def extract_text_from_pdf(file_path: str, use_ocr: bool = True) -> str:
    """Trích xuất text PDF; nếu quá ít ký tự thì fallback OCR (PDF scan)."""
    text = ""
    try:
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text("text") + "\n"
        doc.close()
    except Exception as e:
        print(f"Loi khi doc PDF {file_path}: {e}")

    if use_ocr and needs_ocr_fallback(text):
        try:
            ocr_text = ocr_pdf_pages(file_path)
            if len(ocr_text.strip()) > len(text.strip()):
                text = ocr_text
        except Exception as e:
            print(f"OCR fallback that bai cho {file_path}: {e}")

    return text


def extract_text_from_docx(file_path: str) -> str:
    """Trích xuất nội dung văn bản từ file Word (.docx) kèm theo định dạng cơ bản."""
    text = ""
    try:
        doc = Document(file_path)

        def process_paragraph(para):
            if not para.text.strip():
                return ""
            font_size, font_name, font_color, is_bold, is_italic = None, None, None, False, False
            for run in para.runs:
                if run.text.strip():
                    if run.font.size:
                        font_size = run.font.size.pt
                    if run.font.name:
                        font_name = run.font.name
                    if run.font.color and run.font.color.rgb:
                        font_color = str(run.font.color.rgb)
                    if run.bold:
                        is_bold = True
                    if run.italic:
                        is_italic = True
                    break

            meta_parts = []
            if font_name:
                meta_parts.append(f'font="{font_name}"')
            if font_size:
                meta_parts.append(f'size="{font_size}pt"')
            if font_color and font_color not in ["000000", "None"]:
                meta_parts.append(f'color="#{font_color}"')
            if is_bold:
                meta_parts.append('bold="true"')
            if is_italic:
                meta_parts.append('italic="true"')

            meta_tag = f"[{' '.join(meta_parts)}] " if meta_parts else ""
            return f"{meta_tag}{para.text}\n"

        for para in doc.paragraphs:
            text += process_paragraph(para)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text += process_paragraph(para)
                text += "\n"
    except Exception as e:
        print(f"Loi khi doc DOCX {file_path}: {e}")
    return text


def extract_plain_text_from_docx(file_path: str) -> str:
    """Text thuần cho UI / tóm tắt — không gắn tag metadata."""
    parts = []
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            parts.append(para.text)
    except Exception as e:
        print(f"Loi plain text DOCX {file_path}: {e}")
    return "\n".join(parts)


def _run_to_html(run) -> str:
    text = html.escape(run.text or "")
    if not text:
        return ""
    styles = []
    if run.bold:
        styles.append("font-weight:700")
    if run.italic:
        styles.append("font-style:italic")
    if run.font.size:
        styles.append(f"font-size:{run.font.size.pt}pt")
    if run.font.name:
        styles.append(f"font-family:{run.font.name}")
    if run.font.color and run.font.color.rgb:
        rgb = str(run.font.color.rgb)
        if rgb not in ("000000", "None"):
            styles.append(f"color:#{rgb}")
    if styles:
        return f'<span style="{";".join(styles)}">{text}</span>'
    return text


def _paragraph_align(para) -> str:
    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return "center"
    if para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        return "right"
    if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return "justify"
    return "left"


def docx_to_html(file_path: str) -> str:
    """Chuyển DOCX sang HTML để hiển thị preview (không dùng cho AI)."""
    blocks = ['<div class="docx-preview-inner">']
    try:
        doc = Document(file_path)
        for para_idx, para in enumerate(doc.paragraphs):
            align = _paragraph_align(para)
            inner = "".join(_run_to_html(r) for r in para.runs) or html.escape(para.text)
            if not inner.strip():
                blocks.append(f'<p class="docx-p-empty" data-para-idx="{para_idx}">&nbsp;</p>')
            else:
                blocks.append(
                    f'<p class="docx-p" data-para-idx="{para_idx}" '
                    f'contenteditable="true" style="text-align:{align}">{inner}</p>'
                )
        for table in doc.tables:
            blocks.append('<table class="docx-table">')
            for row in table.rows:
                blocks.append("<tr>")
                for cell in row.cells:
                    cell_html = []
                    for para in cell.paragraphs:
                        inner = "".join(_run_to_html(r) for r in para.runs) or html.escape(para.text)
                        if inner.strip():
                            cell_html.append(f"<p>{inner}</p>")
                    blocks.append(f'<td>{"".join(cell_html) or "&nbsp;"}</td>')
                blocks.append("</tr>")
            blocks.append("</table>")
    except Exception as e:
        print(f"Loi docx_to_html {file_path}: {e}")
        return f'<p class="docx-preview-error">Không thể hiển thị file: {html.escape(str(e))}</p>'
    blocks.append("</div>")
    return "".join(blocks)


def extract_plain_text(file_path: str, use_ocr: bool = True) -> str:
    """Text thuần theo đuôi file (không metadata tag)."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path, use_ocr=use_ocr)
    if ext == ".docx":
        return extract_plain_text_from_docx(file_path)
    raise ValueError(f"Định dạng file không được hỗ trợ: {ext}")


def extract_text(file_path: str, use_ocr: bool = True) -> str:
    """Tự động nhận diện đuôi file để gọi hàm trích xuất tương ứng."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path, use_ocr=use_ocr)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Định dạng file không được hỗ trợ: {ext}")
